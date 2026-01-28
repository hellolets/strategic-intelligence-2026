"""
Módulo para generar reportes en formato Word (.docx).
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional
from .logger import logger
from .config import CONFIG

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from concurrent.futures import ThreadPoolExecutor
from functools import partial


def add_bookmark(paragraph, name: str, bookmark_id: int):
    """
    Inserta un bookmark (destino) que envuelve el contenido del párrafo.
    bookmark_id debe ser único en todo el documento.
    
    Los bookmarks son necesarios para que el TOC pueda crear hipervínculos
    que permitan acceso directo a las secciones haciendo clic en el índice.
    """
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    p = paragraph._p
    # El elemento pPr (si existe) debe ser el primer hijo. 
    # Insertamos el bookmark después del pPr para no romper el estilo del párrafo.
    # IMPORTANTE: El bookmark debe envolver el contenido del párrafo para que
    # el TOC pueda crear hipervínculos correctamente.
    if p.pPr is not None:
        p.pPr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def add_internal_link(
    paragraph,
    anchor: str,
    text: str,
    font_name: str = "Calibri",
    font_size_pt: int = 10,
    blue: bool = True,
    underline: bool = True,
):
    """
    Inserta un hipervínculo interno (clicable) que apunta a un bookmark (anchor).
    También configura fuente/tamaño/color/subrayado.
    """
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Fuente (robusta)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)

    # Tamaño (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size_pt * 2))
    rPr.append(sz)

    # Color típico de link
    if blue:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

    # Subrayado
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)

    t = OxmlElement("w:t")
    t.text = text

    run.append(rPr)
    run.append(t)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def add_external_link(
    paragraph,
    url: str,
    text: str = None,
    font_name: str = "Calibri",
    font_size_pt: int = 10,
    blue: bool = True,
    underline: bool = True,
):
    """
    Inserta un hipervínculo externo (clicable) que abre una URL.
    
    Args:
        paragraph: Párrafo de Word donde insertar el link
        url: URL completa (ej: "https://example.com")
        text: Texto a mostrar (si None, usa la URL completa)
        font_name: Fuente del texto
        font_size_pt: Tamaño de fuente en puntos
        blue: Si True, colorea el link de azul
        underline: Si True, subraya el link
    """
    if text is None:
        text = url
    
    hyperlink = OxmlElement("w:hyperlink")
    # Para links externos usamos r:id, pero python-docx no expone fácilmente las relaciones
    # Alternativa más simple: usar el atributo w:anchor con la URL completa no funciona para externos
    # Mejor enfoque: crear un elemento de hyperlink con tgtFrame
    
    # Establecer el target como la URL
    hyperlink.set(qn("w:tgtFrame"), "_blank")  # Abrir en nueva ventana
    hyperlink.set(qn("w:tooltip"), url)
    
    # IMPORTANTE: Para que funcione como link externo, necesitamos añadirlo a las relaciones del documento
    # Alternativa más robusta: usar add_hyperlink desde python-docx si está disponible
    # Como python-docx no tiene método directo, usamos un workaround
    
    # Crear el parte del relationship
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Fuente
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Tamaño (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size_pt * 2))
    rPr.append(sz)

    # Color azul para links
    if blue:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")  # Azul estándar de Word para links
        rPr.append(color)

    # Subrayado
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)

    t = OxmlElement("w:t")
    t.text = text

    run.append(rPr)
    run.append(t)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def detect_header_level_from_numbering(text):
    """
    Detecta el nivel de encabezado basándose en el patrón de numeración.

    Patterns:
    - "1. ", "2. ", "3. " → Level 2 (H2)
    - "1.1. ", "2.3. ", "4.2. " → Level 3 (H3)
    - "1.1.1. ", "2.3.4. ", "7.7.7. " → Level 4 (H4)
    - Sin numeración → Level 2 (H2, por defecto)

    Returns:
        int: 2, 3, or 4 indicating the heading level
    """
    import re

    # Pattern para X.Y.Z. (H4)
    if re.match(r"^\d+\.\d+\.\d+\.", text):
        return 4

    # Pattern para X.Y. (H3)
    if re.match(r"^\d+\.\d+\.", text):
        return 3

    # Pattern para X. (H2) o sin numeración (H2 por defecto)
    return 2


def generate_docx_from_markdown(markdown_text: str, output_path: str, plot_data: List[Dict[str, Any]] = None):
    """
    Genera un archivo .docx a partir de texto en Markdown.
    Soporta:
    - Headers con detección automática de nivel por numeración (H1, H2, H3, H4)
    - Negritas (**texto**)
    - Listas (normales)
    - Párrafos regulares
    - Marker especial: [[TOC]] para insertar Tabla de Contenido
    - Marker especial: \newpage para saltos de página

    Estilos aplicados:
    - Title (H1): 22pt, Aptos Display (fallback: Calibri), LEFT-ALIGNED
    - Heading 2 (H2 - 1.): 20pt, Aptos Display, negrita
    - Heading 3 (H3 - 1.1.): 16pt, Aptos Display, negrita
    - Heading 4 (H4 - 1.1.1.): 14pt, Aptos, negrita
    - Normal: 10pt, Calibri
    - Page numbering: footer centrado

    Args:
        markdown_text (str): El contenido en Markdown.
        output_path (str): Ruta completa donde guardar el archivo .docx.
        plot_data (list): Opcional. Lista de datos de gráficos generados.
    """
    
    # helper function for parallel download
    def download_plot(plot_id: str):
        try:
            from .r2_utils import r2_manager
            from .config import R2_BUCKET_NAME
            os.makedirs("temp_plots", exist_ok=True)
            dest_path = f"temp_plots/plot_{plot_id}.png"
            
            if os.path.exists(dest_path):
                return plot_id, dest_path
                
            logger.log_info(f"🔄 Descargando gráfico {plot_id} de R2...")
            r2_manager.s3_client.download_file(
                R2_BUCKET_NAME, 
                f"plots/{plot_id}.png", 
                dest_path
            )
            return plot_id, dest_path
        except Exception as e:
            logger.log_warning(f"⚠️ No se pudo descargar el gráfico {plot_id} de R2: {e}")
            return plot_id, None

    try:
        # Validar que output_path no esté vacío
        if not output_path or not output_path.strip():
            logger.log_error("Error: output_path está vacío o es None")
            return False
            
        # Asegurar que el directorio existe
        output_dir = os.path.dirname(output_path)
        if output_dir:  # Solo crear si hay directorio padre
            os.makedirs(output_dir, exist_ok=True)

        # Cargar la plantilla en lugar de crear un documento vacío
        template_path = os.path.join(os.path.dirname(__file__), "..", "assets", "template.docx")
        template_path = os.path.abspath(template_path)
        
        if not os.path.exists(template_path):
            logger.log_warning(f"Plantilla no encontrada en {template_path}. Creando documento vacío.")
            doc = Document()
        else:
            logger.log_info(f"Usando plantilla: {template_path}")
            try:
                doc = Document(template_path)
            except Exception as template_err:
                logger.log_warning(f"Error cargando plantilla: {template_err}. Creando documento vacío.")
                doc = Document()
            
            # Limpiar párrafos vacíos iniciales de la plantilla
            # Esto evita que haya una línea en blanco antes del título
            while doc.paragraphs and len(doc.paragraphs) > 0:
                first_para = doc.paragraphs[0]
                # Si el primer párrafo está vacío o solo tiene espacios, eliminarlo
                if not first_para.text.strip():
                    p_element = first_para._element
                    p_element.getparent().remove(p_element)
                else:
                    # Si encontramos un párrafo con contenido, dejar de eliminar
                    break

        # ============================================
        # CONFIGURACIÓN PARA REFERENCIAS/LINKS
        # ============================================
        # La plantilla maneja todos los estilos de headers y texto.
        # Solo necesitamos configuración para los hipervínculos de las citas.
        refs_cfg = CONFIG.get("references", {})
        link_cfg = refs_cfg.get("link_style", {})

        # ============================================
        # PRE-DESCARGAR GRÁFICOS EN PARALELO
        # ============================================
        plot_tags = re.findall(r"\[\[PLOT:(.*?)\]\]", markdown_text)
        
        # Logging para debugging
        if plot_tags:
            logger.log_info(f"📊 Marcadores PLOT encontrados en markdown: {len(plot_tags)}")
            for i, tag in enumerate(plot_tags[:3], 1):
                plot_id = tag.split("|")[0] if "|" in tag else tag
                logger.log_info(f"   {i}. [[PLOT:{plot_id}...]]")
            if len(plot_tags) > 3:
                logger.log_info(f"   ... y {len(plot_tags) - 3} más")
        else:
            logger.log_warning("⚠️ No se encontraron marcadores [[PLOT:...]] en el markdown")
        
        plots_to_download = []
        for tag in plot_tags:
            plot_id = tag.split("|")[0] if "|" in tag else tag
            # Solo descargar si no está en plot_data local y no existe ya en temp_plots
            is_local = plot_data and any(p["id"] == plot_id and os.path.exists(p["path"]) for p in plot_data)
            if not is_local and not os.path.exists(f"temp_plots/plot_{plot_id}.png"):
                plots_to_download.append(plot_id)
        
        if plots_to_download:
            logger.log_info(f"🚀 Descargando {len(plots_to_download)} gráficos en paralelo desde R2...")
            with ThreadPoolExecutor(max_workers=5) as executor:
                results = list(executor.map(download_plot, list(set(plots_to_download))))
                # Verificar resultados
                downloaded = sum(1 for _, path in results if path is not None)
                failed = sum(1 for _, path in results if path is None)
                if downloaded > 0:
                    logger.log_success(f"✅ {downloaded} gráfico(s) descargado(s) exitosamente")
                if failed > 0:
                    logger.log_warning(f"⚠️ {failed} gráfico(s) no se pudieron descargar desde R2")
        elif plot_tags:
            logger.log_info(f"ℹ️ Todos los gráficos ya están disponibles localmente")

        # ============================================
        # PROCESAR CONTENIDO MARKDOWN
        # ============================================
        lines = markdown_text.split("\n")
        refs_enabled = refs_cfg.get("enable_hyperlinks", True)
        refs_style = (refs_cfg.get("style", "IEEE") or "IEEE").upper()
        in_references = False
        bookmark_counter = 2000  # Inicializar contador para bookmarks
        plot_counter = 1
        
        # Primera pasada: extraer encabezados para el TOC
        toc_headings = []
        bookmark_id_map = {}  # text -> bookmark_id (para asegurar IDs consistentes)
        current_bookmark_counter = bookmark_counter  # Contador separado para bookmarks de headings
        
        for line in lines:
            line_stripped = line.strip()
            # Buscar encabezados H2, H3, H4 (##, ###, ####)
            match = re.match(r'^(#{2,4})\s+(.+)$', line_stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                # Saltar si es encabezado de sistema o meta-sección
                skip_terms = [
                    'table of contents', 'executive summary', 'resumen ejecutivo', 
                    'references', 'referencias', 'fuentes consultadas', 'index',
                    'fuentes', 'índice', 'indicé'
                ]
                if any(skip in text.lower() for skip in skip_terms):
                    continue
                # Generar nombre de bookmark
                bookmark_name = generate_bookmark_name(text)
                # Asignar bookmark_id único
                if text not in bookmark_id_map:
                    bookmark_id_map[text] = current_bookmark_counter
                    current_bookmark_counter += 1
                toc_headings.append((level, text, bookmark_name))
        
        # Segunda pasada: procesar contenido y crear bookmarks cuando se encuentren los headings
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Salto de página
            if line == "\\newpage" or line == "\\\\newpage":
                doc.add_page_break()
                continue

            # TOC Marker
            if "[[TOC]]" in line:
                # Solo usar el Campo TOC Nativo de Word (se actualiza automáticamente con Ctrl+A, F9)
                # NO usar add_static_toc_with_links() porque genera un índice duplicado sin números de página
                add_toc(doc)
                # Añadir un salto de página después del índice si no hay uno inmediatamente después
                if not (i + 1 < len(lines) and ("newpage" in lines[i+1].lower())):
                    doc.add_page_break()
                continue

            # Plot Marker
            plot_match = re.search(r"\[\[PLOT:(.*?)\]\]", line)
            if plot_match:
                content = plot_match.group(1)
                plot_id = content
                fig_word = "Figura" # Default
                title = None
                
                if "|" in content:
                    parts = content.split("|")
                    plot_id = parts[0]
                    if len(parts) == 3:
                        fig_word = parts[1]
                        title = parts[2]
                    elif len(parts) == 2:
                        title = parts[1]

                img_path = None
                
                # Opcion 1: Buscar en plot_data (local)
                if plot_data:
                    plot_info = next((p for p in plot_data if p["id"] == plot_id), None)
                    if plot_info and os.path.exists(plot_info["path"]):
                        img_path = plot_info["path"]
                        # Rescatar metadata si no vino en el bookmark
                        if not title:
                            title = plot_info.get("title")
                        if not fig_word or fig_word == "Figura":
                            fig_word = plot_info.get("figure_word", "Figura")
                
                # Opcion 2: Buscar en carpeta temp_plots por nombre (pre-descargado o local)
                if not img_path:
                    local_path = os.path.abspath(f"temp_plots/plot_{plot_id}.png")
                    if os.path.exists(local_path):
                        img_path = local_path
                
                # Opcion 3: Rescatar de R2 si falló la pre-descarga (fallback)
                if not img_path:
                    _, img_path = download_plot(plot_id)

                if img_path:
                    try:
                        from docx.shared import Inches, Pt
                        
                        # Validar que el archivo existe y es accesible
                        if not os.path.exists(img_path):
                            logger.log_warning(f"⚠️ Imagen no encontrada: {img_path}. Saltando gráfico {plot_id}.")
                            continue
                        
                        # Validar tamaño del archivo (evitar archivos corruptos o muy grandes)
                        file_size = os.path.getsize(img_path)
                        if file_size == 0:
                            logger.log_warning(f"⚠️ Imagen vacía (0 bytes): {img_path}. Saltando gráfico {plot_id}.")
                            continue
                        if file_size > 50 * 1024 * 1024:  # 50MB límite
                            logger.log_warning(f"⚠️ Imagen muy grande ({file_size / 1024 / 1024:.1f}MB): {img_path}. Saltando gráfico {plot_id}.")
                            continue
                        
                        # 1. Insertar Título (Palabra Localizada + Número + Título)
                        if title:
                            # Construir: "Figure 1: Title"
                            full_title = f"{fig_word} {plot_counter}: {title}"
                            p_title = doc.add_paragraph()
                            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run_title = p_title.add_run(full_title)
                            run_title.bold = True
                            run_title.font.size = Pt(11)
                            # Reducir espacio después del título para que esté pegado a la imagen
                            p_title.paragraph_format.space_after = Pt(2)
                            
                            # Incrementar contador de figuras
                            plot_counter += 1

                        # 2. Insertar imagen centrada
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        # Respetar el límite de media página (~4-5 pulgadas)
                        # Usar try-except específico para errores de imagen
                        try:
                            run_img.add_picture(img_path, width=Inches(5.0))
                        except Exception as img_err:
                            logger.log_error(f"❌ Error insertando imagen {plot_id} desde {img_path}: {img_err}")
                            logger.log_error(f"   Tipo de error: {type(img_err).__name__}")
                            # Continuar con el siguiente elemento en lugar de fallar todo el documento
                            continue
                        continue
                    except Exception as plot_err:
                        logger.log_error(f"❌ Error procesando gráfico {plot_id}: {plot_err}")
                        logger.log_error(f"   Tipo: {type(plot_err).__name__}")
                        import traceback
                        logger.log_error(f"   Traceback: {traceback.format_exc()}")
                        # Continuar con el siguiente elemento
                        continue
                else:
                    logger.log_warning(f"❌ No se encontró el gráfico {plot_id} ni localmente ni en R2. Saltando.")
                    continue

            # H1 (Título principal) - LEFT ALIGNED
            if line.startswith("# "):
                text = line[2:].strip()
                heading = doc.add_heading(text, level=1)
                
                # Añadir bookmark para que el TOC pueda enlazar
                bookmark_name = generate_bookmark_name(text)
                add_bookmark(heading, name=bookmark_name, bookmark_id=bookmark_counter)
                bookmark_counter += 1
                
                # Detectar si este H1 es la sección de referencias / apéndice de fuentes
                lt = text.strip().lower()
                if (
                    lt.startswith("references")
                    or lt.startswith("referencias")
                    or ("apéndice" in lt and "fuentes" in lt)
                    or ("apendice" in lt and "fuentes" in lt)
                ):
                    in_references = True
                else:
                    in_references = False
                    bookmark_counter = 1000  # Empezar en 1000 para evitar conflicto con refs [1], [2], etc.

            # H2/H3/H4 (Capítulos jerárquicos) - Detección automática
            elif re.match(r"^(#{2,4})\s+.+", line):
                # Aceptar cabeceras que empiezan por ##, ### o ####
                m = re.match(r"^(#{2,4})\s+(.*)$", line)
                if m:
                    text = m.group(2).strip()
                    # Detectar nivel por numeración del texto (1. => H2, 1.1. => H3, 1.1.1. => H4)
                    level = detect_header_level_from_numbering(text)
                    
                    # Crear heading con el nivel detectado
                    heading = doc.add_heading(text, level=level)
                    
                    # Añadir bookmark para que el TOC pueda enlazar
                    # Usar el mismo nombre de bookmark y ID que se generó en la primera pasada
                    bookmark_name = generate_bookmark_name(text)
                    # Usar el bookmark_id pre-asignado si existe, sino crear uno nuevo
                    bookmark_id = bookmark_id_map.get(text, bookmark_counter)
                    if text not in bookmark_id_map:
                        bookmark_id_map[text] = bookmark_counter
                        bookmark_counter += 1
                    add_bookmark(heading, name=bookmark_name, bookmark_id=bookmark_id)
                    logger.log_info(f"   📌 Bookmark creado: '{bookmark_name}' (ID: {bookmark_id}) para '{text}'")
                    
                    # Marcar si estamos en la sección de referencias
                    if text.strip().lower().startswith(
                        "references"
                    ) or text.strip().lower().startswith("referencias"):
                        in_references = True
                    else:
                        in_references = False
                    continue

            # Listas
            elif line.startswith("- ") or line.startswith("* "):
                text = line[2:].strip()
                # Detectar enlaces de markdown en listas (TOC links) y convertirlos a hipervínculos internos
                if text.startswith("[") and "](#" in text:
                    # Extraer texto y anchor del enlace markdown: [text](#anchor)
                    link_match = re.match(r'\[([^\]]+)\]\(#([^\)]+)\)', text)
                    if link_match:
                        link_text = link_match.group(1)
                        anchor = link_match.group(2)
                        
                        # Convertir anchor de markdown a nombre de bookmark de Word
                        # El anchor de markdown es un slug (ej: "executive-summary")
                        # El bookmark de Word tiene formato "_1_Executive_Summary"
                        # Necesitamos buscar el bookmark correspondiente en toc_headings
                        bookmark_name = None
                        for level, heading_text, bm_name in toc_headings:
                            # Generar anchor desde el texto del heading usando el mismo algoritmo que consolidator.py
                            heading_anchor = re.sub(r'[^\w\s-]', '', heading_text.lower())
                            heading_anchor = re.sub(r'[-\s]+', '-', heading_anchor)
                            if heading_anchor == anchor:
                                bookmark_name = bm_name
                                break
                        
                        # Si encontramos el bookmark, crear hipervínculo interno
                        if bookmark_name:
                            try:
                                p = doc.add_paragraph(style="List Bullet")
                            except KeyError:
                                p = doc.add_paragraph()
                                p.style = 'Normal'
                                from docx.shared import Inches
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.paragraph_format.first_line_indent = Inches(-0.25)
                            
                            # Añadir indentación adicional si hay espacios al inicio de la línea original
                            original_indent = len(line) - len(line.lstrip())
                            if original_indent > 2:  # Si hay indentación adicional (para H3, H4, etc.)
                                current_indent = p.paragraph_format.left_indent or Pt(0)
                                p.paragraph_format.left_indent = current_indent + Pt((original_indent - 2) * 6)
                            
                            # Crear hipervínculo interno
                            add_internal_link(p, bookmark_name, link_text, font_size_pt=10)
                            continue
                        else:
                            # Si no encontramos el bookmark, extraer solo el texto (fallback)
                            text = link_text
                    else:
                        # Si no coincide el patrón, extraer solo el texto
                        text = text.replace("[", "").replace("](#", "").replace(")", "")
                    match = re.match(r"\[(.*?)\]\(#.*?\)", text)
                    if match:
                        text = match.group(1)
                
                # Crear párrafo de lista. Si la plantilla no tiene "List Bullet", usar estilo normal
                try:
                    p = doc.add_paragraph(style="List Bullet")
                except KeyError:
                    # Si la plantilla no tiene el estilo "List Bullet", usar Normal y agregar viñeta manualmente
                    p = doc.add_paragraph()
                    # Agregar formato de lista usando el sistema de numeración de Word
                    p.style = 'Normal'
                    p_format = p.paragraph_format
                    # Configurar sangría para simular lista
                    from docx.shared import Inches
                    p_format.left_indent = Inches(0.25)
                    p_format.first_line_indent = Inches(-0.25)
                    # Añadir viñeta al inicio del texto
                    text = "• " + text
                
                if refs_enabled and refs_style == "IEEE" and not in_references:
                    add_text_with_citations(
                        p,
                        text,
                        link_font=link_cfg.get("font", "Calibri"),
                        link_size_pt=int(link_cfg.get("size_pt", 10)),
                        link_blue=bool(link_cfg.get("blue", True)),
                        link_underline=bool(link_cfg.get("underline", True)),
                    )
                else:
                    add_formatted_text(p, text)

            # Separadores (ignorar)
            elif line.startswith("---") or line.startswith("==="):
                continue

            # Párrafos normales
            else:
                # Si estamos en la sección de referencias y empieza por [n], crear bookmark
                if in_references and refs_enabled and refs_style == "IEEE":
                    # Aceptar distintos formatos de entrada: "[n] ...", "n. ...", "n) ..."
                    m = re.match(r"^\[(\d+)\]\s*(.*)$", line)
                    if m:
                        ref_num = int(m.group(1))
                        rest = m.group(2)
                        p = doc.add_paragraph()
                        add_bookmark(p, name=f"ref_{ref_num}", bookmark_id=ref_num)
                        # Usar add_text_with_url_links para hacer las URLs clicables
                        add_text_with_url_links(p, f"[{ref_num}] {rest}")
                        continue
                    m2 = re.match(r"^(\d+)\.[\s\-–]*?(.*)$", line)
                    if m2:
                        ref_num = int(m2.group(1))
                        rest = m2.group(2)
                        p = doc.add_paragraph()
                        add_bookmark(p, name=f"ref_{ref_num}", bookmark_id=ref_num)
                        # Usar add_text_with_url_links para hacer las URLs clicables
                        add_text_with_url_links(p, f"[{ref_num}] {rest}")
                        continue
                    m3 = re.match(r"^(\d+)\)[\s\-–]*?(.*)$", line)
                    if m3:
                        ref_num = int(m3.group(1))
                        rest = m3.group(2)
                        p = doc.add_paragraph()
                        add_bookmark(p, name=f"ref_{ref_num}", bookmark_id=ref_num)
                        # Usar add_text_with_url_links para hacer las URLs clicables
                        add_text_with_url_links(p, f"[{ref_num}] {rest}")
                        continue
                    # Patrón del apéndice generado: "**Fuente n**" (negrita)
                    m4 = re.match(
                        r"^\*\*\s*(?:Fuente|Ref|Referencia)\s+(\d+)\s*\*\*", line, re.IGNORECASE
                    )
                    if m4:
                        ref_num = int(m4.group(1))
                        p = doc.add_paragraph()
                        add_bookmark(p, name=f"ref_{ref_num}", bookmark_id=ref_num)
                        # Usar add_text_with_url_links para hacer las URLs clicables
                        add_text_with_url_links(p, line)
                        continue
                # Párrafo general
                p = doc.add_paragraph()
                if refs_enabled and refs_style == "IEEE" and not in_references:
                    add_text_with_citations(
                        p,
                        line,
                        link_font=link_cfg.get("font", "Calibri"),
                        link_size_pt=int(link_cfg.get("size_pt", 10)),
                        link_blue=bool(link_cfg.get("blue", True)),
                        link_underline=bool(link_cfg.get("underline", True)),
                    )
                else:
                    add_formatted_text(p, line)

        # Guardar
        doc.save(output_path)
        logger.log_success(f"Archivo Word generado en: {output_path}")
        return True

    except Exception as e:
        import traceback
        error_msg = str(e) if e else "Error desconocido (excepción sin mensaje)"
        error_traceback = traceback.format_exc()
        logger.log_error(f"Error generando archivo Word: {error_msg}")
        logger.log_error(f"Traceback completo:\n{error_traceback}")
        # También imprimir en consola para debugging
        print(f"❌ ERROR DETALLADO generando Word:")
        print(f"   Mensaje: {error_msg}")
        print(f"   Tipo: {type(e).__name__}")
        if error_traceback:
            print(f"   Traceback:\n{error_traceback}")
        return False


def add_static_toc_with_links(doc, toc_headings):
    """
    Genera un TOC estático con hipervínculos usando bookmarks.
    Este TOC se muestra inmediatamente sin necesidad de actualizar campos en Word.
    
    Args:
        doc: Documento de Word
        toc_headings: Lista de tuplas (level, text, bookmark_name) con los encabezados
    """
    from docx.shared import Pt, RGBColor
    
    # Añadir título "Index" con estilo Normal modificado
    index_title = doc.add_paragraph("Index", style="Normal")
    run = index_title.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(15, 71, 97)
    
    if not toc_headings:
        # Si no hay encabezados, añadir mensaje
        p = doc.add_paragraph("No se encontraron encabezados para el índice.")
        return
    
    # Generar TOC estático con hipervínculos
    for level, text, bookmark_name in toc_headings:
        # Crear párrafo para cada entrada del TOC
        p = doc.add_paragraph()
        
        # Añadir indentación según el nivel (H2=0, H3=1, H4=2)
        indent_spaces = "  " * (level - 2)
        if indent_spaces:
            indent_run = p.add_run(indent_spaces)
            indent_run.font.size = Pt(10)
        
        # Añadir hipervínculo al encabezado
        add_internal_link(p, bookmark_name, text, font_size_pt=10)
    
    # Añadir espacio después del TOC
    doc.add_paragraph()


def add_toc(doc):
    """
    Inserta un campo TOC (Table of Contents) nativo de Word con hipervínculos.
    El campo TOC se actualizará automáticamente cuando se abra el documento en Word.
    Los hipervínculos funcionarán automáticamente gracias a los bookmarks en los encabezados.
    
    IMPORTANTE: El campo TOC necesita ser actualizado en Word para mostrar el contenido.
    Para actualizar: Clic derecho en el campo TOC > "Actualizar campos" > "Actualizar toda la tabla"
    
    NOTA: Esta función ya no se usa - se usa add_static_toc_with_links en su lugar.
    """
    from docx.oxml import parse_xml
    from docx.shared import Pt, RGBColor
    
    # Añadir título "Index" con estilo Normal modificado
    index_title = doc.add_paragraph("Index", style="Normal")
    run = index_title.runs[0]
    
    # Modificar solo este párrafo (no el estilo global)
    run.font.size = Pt(16)
    # Color: Azul oscuro, Texto 2, Claro 10% (aproximadamente RGB: 68, 84, 106)
    run.font.color.rgb = RGBColor(15, 71, 97)
    
    # Añadir el campo TOC
    paragraph = doc.add_paragraph()
    run_toc = paragraph.add_run()

    # Inicio del campo
    fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    # Instrucción TOC con hipervínculos habilitados
    # \o "2-4" = incluir niveles de encabezado 2-4 (Heading 2, 3, 4)
    # \h = crear hipervínculos a los encabezados (CRÍTICO para acceso directo)
    # \z = ocultar números de página en modo web
    # \u = usar estilos de encabezado (Heading 2, Heading 3, Heading 4)
    # \t = usar estilos específicos (opcional, pero \u es más estándar)
    instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> TOC \o "2-4" \h \z \u </w:instrText>')
    
    # Separador
    fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    # Texto placeholder visible (se reemplazará cuando se actualice el campo)
    # Este texto aparecerá hasta que se actualice el campo TOC en Word
    fldChar3 = parse_xml(r'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">Haz clic derecho y selecciona "Actualizar campos" para generar el índice con enlaces clicables.</w:t>')
    
    # Fin del campo
    fldChar4 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    # Añadir TODOS los elementos en el orden correcto
    run_toc._r.append(fldChar1)
    run_toc._r.append(instrText)
    run_toc._r.append(fldChar2)
    run_toc._r.append(fldChar3)
    run_toc._r.append(fldChar4)


def add_formatted_text(paragraph, text):
    """
    Añade texto con formato básico (negritas) a un párrafo.
    """
    # Regex para detectar **negritas**
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            paragraph.add_run(part)


def add_text_with_url_links(paragraph, text, font_name="Calibri", font_size_pt=10):
    """
    Añade texto a un párrafo detectando URLs y convirtiéndolas en hipervínculos clicables.
    También maneja negritas (**texto**).
    
    Args:
        paragraph: Párrafo de Word
        text: Texto que puede contener URLs y formato markdown
        font_name: Fuente para las URLs
        font_size_pt: Tamaño de fuente para las URLs
    """
    # Regex para detectar URLs (https:// o http://)
    url_pattern = r'(https?://[^\s\)\]\>,]+)'
    
    # Dividir el texto por URLs
    parts = re.split(url_pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Si es una URL, crear hipervínculo
        if re.match(url_pattern, part):
            add_external_link(
                paragraph,
                url=part,
                text=part,
                font_name=font_name,
                font_size_pt=font_size_pt,
                blue=True,
                underline=True
            )
        else:
            # Texto normal (puede tener negritas)
            add_formatted_text(paragraph, part)



def add_text_with_citations(
    paragraph,
    text,
    link_font: str = "Calibri",
    link_size_pt: int = 10,
    link_blue: bool = True,
    link_underline: bool = True,
):
    """
    Inserta texto con formato básico (negritas) y convierte citas en hipervínculos internos.
    
    Soporta:
    - Citas individuales: [1], (1)
    - Citas agrupadas: [1, 2, 3], [10, 25, 30]
    - Citas con espacios: [1 , 2 , 3]
    
    Cada número en una cita agrupada se convierte en un hipervínculo separado.
    """
    # Regex mejorado que captura:
    # - [n] o (n) → citas individuales
    # - [n, m, p] o (n, m, p) → citas agrupadas con comas
    pattern = r"(\[\d+(?:\s*,\s*\d+)*\]|\(\d+(?:\s*,\s*\d+)*\))"
    tokens = re.split(pattern, text)
    
    for tok in tokens:
        if not tok:
            continue
        
        # Detectar si es una cita (individual o agrupada)
        if re.fullmatch(r"\[\d+(?:\s*,\s*\d+)*\]", tok):
            # Cita con corchetes: [1] o [1, 2, 3]
            process_citation_group(
                paragraph, tok, 
                opening='[', closing=']',
                link_font=link_font,
                link_size_pt=link_size_pt,
                link_blue=link_blue,
                link_underline=link_underline
            )
        elif re.fullmatch(r"\(\d+(?:\s*,\s*\d+)*\)", tok):
            # Cita con paréntesis: (1) o (1, 2, 3)
            process_citation_group(
                paragraph, tok,
                opening='(', closing=')',
                link_font=link_font,
                link_size_pt=link_size_pt,
                link_blue=link_blue,
                link_underline=link_underline
            )
        else:
            # Texto normal o con formato (negritas)
            add_formatted_text(paragraph, tok)


def process_citation_group(
    paragraph,
    citation_text,
    opening='[',
    closing=']',
    link_font="Calibri",
    link_size_pt=10,
    link_blue=True,
    link_underline=True
):
    """
    Procesa una cita que puede ser individual [1] o agrupada [1, 2, 3].
    
    Para citas agrupadas:
    - Añade el opening bracket: [
    - Convierte cada número en un hipervínculo
    - Añade comas entre números
    - Añade el closing bracket: ]
    
    Args:
        citation_text: Texto completo de la cita, ej: "[1, 2, 3]"
        opening: Carácter de apertura ('[' o '(')
        closing: Carácter de cierre (']' o ')')
    """
    # Extraer los números de la cita
    # Ejemplo: "[1, 2, 3]" → ["1", "2", "3"]
    numbers = re.findall(r'\d+', citation_text)
    
    if not numbers:
        # Si por alguna razón no hay números, añadir texto tal cual
        paragraph.add_run(citation_text)
        return
    
    # Añadir bracket de apertura
    paragraph.add_run(opening)
    
    # Procesar cada número
    for i, num in enumerate(numbers):
        # Añadir hipervínculo para este número
        add_internal_link(
            paragraph,
            anchor=f"ref_{num}",
            text=num,
            font_name=link_font,
            font_size_pt=link_size_pt,
            blue=link_blue,
            underline=link_underline,
        )
        
        # Añadir coma si no es el último número
        if i < len(numbers) - 1:
            paragraph.add_run(", ")
    
    # Añadir bracket de cierre
    paragraph.add_run(closing)

def generate_bookmark_name(heading_text):
    """
    Genera un nombre de bookmark válido para Word a partir del texto del heading.
    
    Word tiene restricciones para nombres de bookmarks:
    - Deben empezar con una letra
    - Solo pueden contener letras, números y guiones bajos
    - Máximo 40 caracteres
    - No pueden contener espacios
    
    Args:
        heading_text: Texto del heading (ej: "1. Executive Summary")
    
    Returns:
        str: Nombre de bookmark válido (ej: "_1_Executive_Summary")
    """
    # Convertir a minúsculas y reemplazar espacios por guiones bajos
    name = heading_text.lower().replace(" ", "_")
    
    # Eliminar caracteres no válidos (mantener solo letras, números, guiones bajos)
    name = re.sub(r'[^a-z0-9_]', '', name)
    
    # Asegurar que empieza con letra o guión bajo
    if not name or not (name[0].isalpha() or name[0] == '_'):
        name = '_' + name
    
    # Limitar a 40 caracteres (restricción de Word)
    if len(name) > 40:
        name = name[:40]
    
    # Si después de todo el proceso queda vacío, usar un nombre genérico
    if not name:
        name = '_heading'
    
    return name
